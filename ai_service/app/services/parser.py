"""
Document text extraction for PDF, DOCX, TXT and MD uploads.

Hardened from v1: encoding errors on text files are tolerated rather than
raised (a single stray byte in an otherwise fine document should not fail the
whole ingest), DOCX table content is extracted instead of silently dropped,
and a document that yields no text is reported as such rather than being
indexed as an empty record.
"""
from __future__ import annotations

import logging
import os
from typing import List

from app.core.observability import safe_extra

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md"})


def _parse_pdf(file_path: str) -> str:
    import pypdf

    parts: List[str] = []
    with open(file_path, "rb") as handle:
        reader = pypdf.PdfReader(handle)
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text()
            except Exception as exc:  # noqa: BLE001
                # One unparseable page (odd fonts, damaged xref) shouldn't
                # discard the rest of an otherwise readable document.
                logger.warning("Skipping unreadable PDF page %d: %s", page_number, exc)
                continue
            if text:
                parts.append(text)
    return "\n".join(parts)


def _parse_docx(file_path: str) -> str:
    import docx

    document = docx.Document(file_path)
    parts: List[str] = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables frequently hold the specification data that makes a product
    # document worth indexing; v1 dropped them entirely.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _parse_text(file_path: str) -> str:
    # errors="replace" keeps a mostly-valid file usable instead of failing the
    # entire ingest over a handful of undecodable bytes.
    with open(file_path, "r", encoding="utf-8", errors="replace") as handle:
        return handle.read()


def parse_file(file_path: str, filename: str) -> str:
    """
    Extracts plain text from a supported document.

    Raises ValueError for unsupported types, unreadable files, or documents
    that contain no extractable text (a scanned PDF with no OCR layer, for
    example) — the caller turns that into a 400 rather than indexing nothing.
    """
    extension = os.path.splitext(filename)[1].lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file extension '{extension}'. Supported: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}."
        )

    try:
        if extension == ".pdf":
            text = _parse_pdf(file_path)
        elif extension == ".docx":
            text = _parse_docx(file_path)
        else:
            text = _parse_text(file_path)
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to parse %s: %s", filename, exc)
        raise ValueError(f"Could not read '{filename}': {exc}") from exc

    text = text.strip()
    if not text:
        raise ValueError(
            f"No text could be extracted from '{filename}'. If this is a scanned "
            f"document, it needs OCR before it can be indexed."
        )

    logger.info(
        "Parsed document",
        extra=safe_extra(filename=filename, extracted_chars=len(text)),
    )
    return text
